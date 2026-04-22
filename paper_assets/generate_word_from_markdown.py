from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"C:\Users\89657\Desktop\yolov8")
MD_PATH = ROOT / "交通目标检测论文初稿.md"
OUT_PATH = ROOT / "面向道路场景的改进YOLOv8s交通目标检测系统设计与验证_无图片占位版.docx"


FIG_TABLE_MAP = {
    "图 1": "图3.1",
    "图 2": "图4.1",
    "图 3": "图4.2",
    "图 4": "图4.3",
    "图 5": "图5.1",
    "表 1": "表3.1",
    "表 2": "表4.1",
    "表 3": "表4.2",
}

IMAGE_INSERTION_HINTS = {
    "model_architecture.png": "请在此处插入：图3.1 改进 YOLOv8s 结构示意图",
    "training_curves_full.png": "请在此处插入：图4.1 full 变体在 BDD100K 数据集上的训练收敛曲线",
    "confusion_matrix_normalized.png": "请在此处插入：图4.2 full 变体在验证集上的归一化混淆矩阵",
    "detection_examples.png": "请在此处插入：图4.3 交通场景检测结果可视化示例",
    "system_architecture.png": "请在此处插入：图5.1 TrafficDetect 系统总体架构图",
}

CHINESE_NUM = {
    "1": "一",
    "2": "二",
    "3": "三",
    "4": "四",
    "5": "五",
    "6": "六",
}

TOC_PAGE_HINTS = {
    "1": "1",
    "2": "2",
    "3": "5",
    "4": "7",
    "5": "11",
    "6": "13",
    "参考文献": "14",
    "致谢": "16",
}


def replace_numbering(text: str) -> str:
    return re.sub(
        r"(图 [1-5]|表 [1-3])(?!\.)",
        lambda match: FIG_TABLE_MAP.get(match.group(1), match.group(1)),
        text,
    )


def format_heading_for_body(heading: str) -> str:
    match = re.match(r"^(\d+)\s+(.+)$", heading)
    if match and match.group(1) in CHINESE_NUM:
        return f"第{CHINESE_NUM[match.group(1)]}章 {match.group(2)}"
    return heading


def format_heading_for_toc(heading: str) -> str:
    return format_heading_for_body(heading)


def set_east_asia_font(run, east_asia: str, ascii_font: str | None = None):
    if ascii_font:
        run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_format(paragraph, first_line: bool = False, center: bool = False):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_line:
        pf.first_line_indent = Pt(24)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_runs_with_citations(paragraph, text: str, *, size=12, east_asia="宋体", ascii_font="Times New Roman", bold=False):
    text = replace_numbering(text)
    pattern = re.compile(r"(\[(?:\d+|[0-9]+[-,，][0-9,\-，]+)\])")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            run.font.size = Pt(size)
            run.bold = bold
            set_east_asia_font(run, east_asia, ascii_font)
        cite = paragraph.add_run(match.group(1).replace("，", ","))
        cite.font.size = Pt(size)
        cite.font.superscript = True
        cite.bold = bold
        set_east_asia_font(cite, east_asia, ascii_font)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(size)
        run.bold = bold
        set_east_asia_font(run, east_asia, ascii_font)


def add_plain_paragraph(doc, text: str, *, first_line=True):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line)
    add_runs_with_citations(p, text)
    return p


def add_center_text(doc, text: str, *, size=12, east_asia="宋体", ascii_font="Times New Roman", bold=False):
    p = doc.add_paragraph()
    set_paragraph_format(p, center=True)
    p.paragraph_format.line_spacing = Pt(max(18, int(size * 1.6)))
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    set_east_asia_font(run, east_asia, ascii_font)
    return p


def add_heading(doc, text: str, level: int):
    text = replace_numbering(format_heading_for_body(text))
    p = doc.add_paragraph()
    if level == 1:
        p.style = "PaperHeading1"
    elif level == 2:
        p.style = "PaperHeading2"
    else:
        p.style = "PaperHeading3"
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        p.paragraph_format.line_spacing = Pt(26)
    elif level == 2:
        run.font.size = Pt(15)
        p.paragraph_format.line_spacing = Pt(24)
    else:
        run.font.size = Pt(14)
        p.paragraph_format.line_spacing = Pt(22)
    set_east_asia_font(run, "黑体", "Times New Roman")
    return p


def set_cell_text(cell, text: str, *, bold=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p)
    run = p.add_run(replace_numbering(text))
    run.bold = bold
    run.font.size = Pt(10.5)
    set_east_asia_font(run, "宋体", "Times New Roman")


def set_cell_borders(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if edge_data is None:
            if element is not None:
                borders.remove(element)
            continue
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def format_three_line_table(table):
    rows = table.rows
    for r_idx, row in enumerate(rows):
        for cell in row.cells:
            set_cell_borders(cell)
            if r_idx == 0:
                set_cell_borders(
                    cell,
                    top={"val": "single", "sz": "12", "color": "000000"},
                    bottom={"val": "single", "sz": "4", "color": "000000"},
                )
            elif r_idx == len(rows) - 1:
                set_cell_borders(cell, bottom={"val": "single", "sz": "12", "color": "000000"})


def add_table(doc, caption: str, rows: list[list[str]]):
    add_center_text(doc, replace_numbering(caption), size=10.5, east_asia="黑体", bold=True)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            set_cell_text(table.cell(i, j), value, bold=(i == 0))
    format_three_line_table(table)
    return table


def add_image(doc, image_path: Path, caption: str):
    caption = replace_numbering(caption)
    hint = IMAGE_INSERTION_HINTS.get(image_path.name, f"请在此处插入：{caption}")
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    row = table.rows[0]
    row.height = Cm(6.2)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_borders(
        cell,
        top={"val": "single", "sz": "8", "color": "999999"},
        bottom={"val": "single", "sz": "8", "color": "999999"},
        left={"val": "single", "sz": "8", "color": "999999"},
        right={"val": "single", "sz": "8", "color": "999999"},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, first_line=False)
    run = p.add_run(hint)
    run.font.size = Pt(12)
    run.bold = True
    set_east_asia_font(run, "宋体", "Times New Roman")
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, first_line=False)
    run = p.add_run(f"本地图片文件：论文插图_手动插入\\{image_path.name}")
    run.font.size = Pt(10.5)
    set_east_asia_font(run, "宋体", "Times New Roman")
    add_center_text(doc, replace_numbering(caption), size=10.5, east_asia="黑体", bold=True)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.size = Pt(10.5)
    set_east_asia_font(run, "宋体", "Times New Roman")


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "右键更新域以生成目录"
    fld_sep.append(text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_manual_toc(doc, sections):
    add_center_text(doc, "目 录", size=14, east_asia="宋体", bold=False)
    toc_items = []
    for level, heading, _ in sections:
        if heading in {"摘要", "Abstract"}:
            continue
        if heading == "参考文献":
            toc_items.append((1, heading, TOC_PAGE_HINTS["参考文献"]))
            continue
        if re.match(r"^\d+\s+", heading):
            chapter_no = re.match(r"^(\d+)\s+", heading).group(1)
            toc_items.append((1, format_heading_for_toc(heading), TOC_PAGE_HINTS.get(chapter_no, "")))
        elif re.match(r"^\d+\.\d+\s+", heading):
            chapter_no = heading.split(".", 1)[0]
            toc_items.append((2, heading, TOC_PAGE_HINTS.get(chapter_no, "")))
    toc_items.append((1, "致谢", TOC_PAGE_HINTS["致谢"]))
    for level, heading, page in toc_items:
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=False)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.2))
        if level == 2:
            p.paragraph_format.left_indent = Pt(24)
        run = p.add_run(f"{heading}\t{page}")
        run.font.size = Pt(12)
        set_east_asia_font(run, "宋体", "Times New Roman")


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.75)
    section.right_margin = Cm(2.5)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for style_name, size, outline in [
        ("PaperHeading1", 16, 0),
        ("PaperHeading2", 15, 1),
        ("PaperHeading3", 14, 2),
    ]:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        ppr = style._element.get_or_add_pPr()
        outline_level = OxmlElement("w:outlineLvl")
        outline_level.set(qn("w:val"), str(outline))
        ppr.append(outline_level)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(18)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)
    return doc


def parse_markdown():
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    title = lines[0].lstrip("# ").strip()
    english_title = lines[2].lstrip("# ").strip()
    author = "待填写"
    unit = "待填写"
    sections: list[tuple[str, str, list[str]]] = []
    current_level = None
    current_title = None
    current_lines: list[str] = []
    for line in lines[4:]:
        if line.startswith("## "):
            if current_title is not None:
                sections.append((current_level, current_title, current_lines))
            current_level = "h1"
            current_title = line[3:].strip()
            current_lines = []
        elif line.startswith("### "):
            if current_title is not None:
                sections.append((current_level, current_title, current_lines))
            current_level = "h2"
            current_title = line[4:].strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_title is not None:
        sections.append((current_level, current_title, current_lines))
    return title, english_title, author, unit, sections


def collect_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-+:?", cell) for cell in row):
            rows.append(row)
        i += 1
    return rows, i


def render_section_content(doc, lines: list[str], *, in_references=False):
    i = 0
    pending_caption = None
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("**") and line.endswith("**"):
            pending_caption = line.strip("*")
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines) and lines[i].strip().startswith("|"):
                rows, i = collect_table(lines, i)
                add_table(doc, pending_caption, rows)
                pending_caption = None
            continue
        if line.startswith("|"):
            rows, i = collect_table(lines, i)
            add_table(doc, pending_caption or "表", rows)
            pending_caption = None
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            caption = lines[i + 2].strip() if i + 2 < len(lines) else image_match.group(1)
            add_image(doc, Path(image_match.group(2)), caption)
            i += 3
            continue
        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False)
            p.paragraph_format.left_indent = Pt(24)
            add_runs_with_citations(p, re.sub(r"^\d+\.\s+", "", line))
            i += 1
            continue
        if in_references and re.match(r"^\[\d+\]", line):
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(21)
            p.paragraph_format.hanging_indent = Pt(21)
            run = p.add_run(line)
            run.font.size = Pt(10.5)
            if re.search(r"^[\[\]0-9,\.\sA-Za-z]", line):
                set_east_asia_font(run, "宋体", "Times New Roman")
            else:
                set_east_asia_font(run, "宋体", "Times New Roman")
            i += 1
            continue
        add_plain_paragraph(doc, line)
        i += 1


def add_cover(doc, title: str, english_title: str, author: str, unit: str):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False)
    add_runs_with_citations(p, "分类号： TP391.4                          单位代码：      11", size=10, bold=True)
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False)
    add_runs_with_citations(p, "密  级：   一般                           学   号：      待填写", size=10, bold=True)
    for _ in range(2):
        doc.add_paragraph()
    add_center_text(doc, "本科毕业论文（设计）", size=20, east_asia="黑体", bold=True)
    for label, value in [
        ("题    目", title),
        ("学    院", "数学与计算机科学学院"),
        ("专    业", "软件工程"),
        ("班    级", "待填写"),
        ("姓    名", author),
        ("指导教师", "待填写"),
        ("职    称", "待填写"),
        ("企业教师", "待填写"),
        ("企业名称", "待填写"),
        ("答辩日期", "二〇二六年六月"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_format(p)
        p.paragraph_format.first_line_indent = Pt(96)
        run = p.add_run(f"{label}：{value}")
        run.font.size = Pt(12)
        run.bold = True
        set_east_asia_font(run, "宋体", "Times New Roman")
    doc.add_page_break()


def add_declaration(doc):
    add_center_text(doc, "延安大学学士学位论文原创性声明", size=14, east_asia="宋体", bold=True)
    add_plain_paragraph(
        doc,
        "本人郑重声明：所呈交的学位论文，是本人在导师的指导下，独立进行研究工作所取得的成果。除文中已经注明引用的内容外，本论文不含任何其他个人或集体已经发表或撰写过的作品成果。对本文的研究做出重要贡献的个人和集体，均已在文中以明确方式标明。本人完全意识到本声明的法律结果由本人承担。",
    )
    add_plain_paragraph(doc, "作者签名：                 日期：", first_line=False)
    doc.add_page_break()
    add_center_text(doc, "关于论文使用授权的说明", size=14, east_asia="宋体", bold=True)
    add_plain_paragraph(
        doc,
        "学位论文作者完全了解延安大学有关保留和使用学位论文的规定，即：本科生在校攻读学士学位期间论文工作的知识产权单位属延安大学，学生公开发表需经指导教师同意。学校有权保留并向国家有关部门或机构送交论文的复印件，允许学位论文被查阅和借阅；学校可以公布学位论文的全部或部分内容，可以允许采用影印、缩印或其它复制手段保存、汇编学位论文。",
    )
    add_plain_paragraph(doc, "保密论文注释：本学位论文属于保密范围，在 2 年解密后适用本授权书。非保密论文注释：本学位论文不属于保密范围，适用本授权书。", first_line=True)
    add_plain_paragraph(doc, "作者签名：                   日期：", first_line=False)
    add_plain_paragraph(doc, "导师签名：                   日期：", first_line=False)
    doc.add_paragraph()


def render_chinese_abstract(doc, title: str, lines: list[str]):
    add_center_text(doc, title, size=14, east_asia="黑体", bold=True)
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("**关键词**") or line.startswith("**关键词"):
            text = re.sub(r"^\*\*关键词\*\*[:：]?", "关键词：", line)
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False)
            p.paragraph_format.first_line_indent = Pt(24)
            label, content = text.split("：", 1)
            run = p.add_run(label + "：")
            run.font.size = Pt(12)
            run.bold = True
            set_east_asia_font(run, "黑体", "Times New Roman")
            run = p.add_run(content.rstrip("。；;"))
            run.font.size = Pt(12)
            set_east_asia_font(run, "宋体", "Times New Roman")
        else:
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False)
            p.paragraph_format.first_line_indent = Pt(24)
            run = p.add_run("摘  要：")
            run.font.size = Pt(12)
            run.bold = True
            set_east_asia_font(run, "黑体", "Times New Roman")
            add_runs_with_citations(p, line, size=12)


def render_english_abstract(doc, english_title: str, lines: list[str]):
    add_center_text(doc, english_title, size=14, east_asia="宋体", ascii_font="Times New Roman", bold=True)
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("**Keywords**") or line.startswith("**Keywords"):
            text = re.sub(r"^\*\*Keywords\*\*[:：]?", "Key words: ", line)
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False)
            run = p.add_run(text)
            run.font.size = Pt(12)
            set_east_asia_font(run, "宋体", "Times New Roman")
            if run.text.startswith("Key words"):
                run.bold = False
        else:
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False)
            run = p.add_run("Abstract: ")
            run.font.size = Pt(12)
            run.bold = True
            set_east_asia_font(run, "宋体", "Times New Roman")
            run = p.add_run(line)
            run.font.size = Pt(12)
            set_east_asia_font(run, "宋体", "Times New Roman")


def main():
    title, english_title, author, unit, sections = parse_markdown()
    doc = setup_document()
    add_cover(doc, title, english_title, author, unit)
    add_declaration(doc)

    add_manual_toc(doc, sections)
    doc.add_page_break()

    for level, heading, lines in sections:
        if heading == english_title:
            continue
        if heading == "摘要":
            render_chinese_abstract(doc, title, lines)
            doc.add_page_break()
        elif heading == "Abstract":
            render_english_abstract(doc, english_title, lines)
            doc.add_page_break()
        elif heading == "参考文献":
            doc.add_page_break()
            add_heading(doc, "参考文献", 1)
            render_section_content(doc, lines, in_references=True)
        elif level == "h1":
            if re.match(r"^\d+\s+", heading) and heading.startswith("1 "):
                doc.add_page_break()
            add_heading(doc, heading, 1)
            render_section_content(doc, lines)
        else:
            add_heading(doc, heading, 2)
            render_section_content(doc, lines)

    doc.add_page_break()
    add_heading(doc, "致谢", 1)
    add_plain_paragraph(doc, "感谢指导教师在论文选题、系统实现和论文写作过程中给予的指导与帮助；感谢同学和朋友在资料整理、实验运行和系统测试过程中提供的支持。", first_line=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
